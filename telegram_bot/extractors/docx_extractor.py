"""
DOCX text extraction: paragraphs and tables.
"""
from __future__ import annotations

import io
from typing import Optional

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None  # type: ignore


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from .docx: all paragraphs and table cells."""
    if Document is None:
        raise RuntimeError("Install python-docx: pip install python-docx")

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append((cell.text or "").strip())
            line = " | ".join(row_text)
            if line.strip():
                parts.append(line)
        if parts and parts[-1].strip():
            parts.append("")  # separator after table

    return "\n".join(parts).strip()
